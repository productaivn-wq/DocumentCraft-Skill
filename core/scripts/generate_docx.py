"""DOCX generation engine for DocumentCraft.

Uses python-docx + htmldocx to render Markdown into styled Word documents.
Supports parallel multi-theme generation via threading.
"""

import os
import threading
from typing import Callable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from htmldocx import HtmlToDocx

from lib.md_parser import read_markdown
from lib.themes_docx import DOCX_THEMES


def _add_page_number(run) -> None:
    """Inject Word XML fields into a run to display the current page number."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _apply_table_borders(doc: Document) -> None:
    """Post-process: apply 'Table Grid' style to all tables for visible borders."""
    for table in doc.tables:
        try:
            table.style = "Table Grid"
        except KeyError:
            pass


def inject_cover_page(doc: Document, metadata: dict) -> None:
    """Inject a cover page into the Word document using parsed YAML metadata."""
    if not metadata or not metadata.get("title"):
        return
        
    title = metadata.get("title", "")
    subtitle = metadata.get("subtitle", "")
    brand = metadata.get("brand", "")
    date = metadata.get("date", "")
    
    # Add spacing
    doc.add_paragraph("\n\n\n\n")
    
    if brand:
        p = doc.add_paragraph(brand)
        p.style = "Subtitle"
        
    p = doc.add_paragraph(title)
    p.style = "Title"
    
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.style = "Subtitle"
        
    if date:
        doc.add_paragraph("\n\n")
        p = doc.add_paragraph(date)
        p.style = "Normal"
        
    doc.add_page_break()


def generate_single(html_content: str, metadata: dict, theme_name: str, output_path: str, brand_color: str = None) -> None:
    """Generate a single DOCX with the specified theme.

    Args:
        html_content: HTML string to inject.
        metadata: Parsed YAML frontmatter for cover page.
        theme_name: Key from DOCX_THEMES (classic, minimalist, vibrant).
        output_path: Destination file path.
        brand_color: Optional hex string.
    """
    if theme_name not in DOCX_THEMES:
        raise ValueError(f"Unknown DOCX theme '{theme_name}'. Available: {list(DOCX_THEMES.keys())}")

    print(f"  [DOCX] Generating {os.path.basename(output_path)} ({theme_name})...")
    doc = Document()

    # Apply theme (mutates Word Styles XML)
    theme_func = DOCX_THEMES[theme_name]
    theme_func(doc, brand_color_hex=brand_color)

    # Inject cover page
    inject_cover_page(doc, metadata)

    # Inject HTML content
    parser = HtmlToDocx()
    parser.add_html_to_document(html_content, doc)

    # Ensure all tables have visible borders
    _apply_table_borders(doc)

    # Add page numbers to footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    _add_page_number(run)

    doc.save(output_path)
    print(f"  [DOCX] Saved: {output_path}")


def generate_all(md_file: str, output_dir: str, themes: Optional[list[str]] = None,
                 parallel: bool = True, brand_color: str = None) -> list[str]:
    """Generate DOCX files for multiple themes.

    Args:
        md_file: Path to the source Markdown file.
        output_dir: Directory to write output DOCX files.
        themes: List of theme names, or None for all themes.
        parallel: If True, render themes concurrently via threads.
        brand_color: Optional hex string to override theme colors.

    Returns:
        List of output file paths.
    """
    os.makedirs(output_dir, exist_ok=True)
    html_content, metadata = read_markdown(md_file)
    theme_list = themes or list(DOCX_THEMES.keys())
    base = os.path.splitext(os.path.basename(md_file))[0]

    paths = []
    threads = []

    for theme_name in theme_list:
        out_path = os.path.join(output_dir, f"{base}_docx_{theme_name}.docx")
        paths.append(out_path)
        if parallel:
            t = threading.Thread(target=generate_single, args=(html_content, metadata, theme_name, out_path, brand_color))
            threads.append(t)
        else:
            generate_single(html_content, metadata, theme_name, out_path, brand_color)

    if parallel and threads:
        for t in threads:
            t.start()
        for t in threads:
            t.join()

    return paths


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python generate_docx.py <markdown_file> [output_dir] [brand_color]")
        sys.exit(1)
    md = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "./outputs"
    color = sys.argv[3] if len(sys.argv) > 3 else None
    results = generate_all(md, out, brand_color=color)
    print(f"\nGenerated {len(results)} DOCX files.")
